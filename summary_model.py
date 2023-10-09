from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain.document_loaders import YoutubeLoader

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

from utils import llm_choice, handle_long_text, extract_text
from docx import Document
import io
import streamlit as st

prompt = PromptTemplate(
    input_variables=["lang", "matter", "text", "features"],
    template="Write a detailed easy to understand summary {lang} of the following: {matter}. {features}: {text}",
)


def summarization_chain(llm_name: str, key: str, files_list: list, lang: str, features: str) -> str:
    """
    Obtains the summary of all files.
    Firstly, it extracts the text from each file, then selects the model and retrieves the LLM and the context,
    which is the number of tokens the model operates with.
    If the extracted text has a number of tokens greater than or equal to what is allowed by the model,
    the 'handle_long_text' function will be used to handle large files, where the text needs to be processed in parts.
    This function returns the summary of that large file, which, along with the file's name,
    is added to the summaries list. In case the text tokens are fewer than the allowed limit,
    the summary of that file is obtained, and similarly, the file's name with its summary is added to the
    'summaries' list. Finally, the summaries of all files in the 'summaries' list are concatenated and returned.

    :param llm_name: Name of the desired LLM.
    :param key: OpenAI API key.
    :param files_list: List of uploaded files
    :param lang: Chosen language
    :param features: Additional features for the summary
    :return: Summaries of all files
    """
    matter = "text"
    summaries = []
    for file in files_list:
        text = extract_text(file)
        llm, context_length = llm_choice(llm_name, key, "sum")
        if llm.get_num_tokens(text) >= context_length:
            long_response = handle_long_text(llm, context_length, text, lang, matter, features)
            summaries.append([file.name, long_response])
        else:
            chain = LLMChain(llm=llm, prompt=prompt)
            summaries.append([file.name, chain.run({'lang': lang, 'matter': matter,
                                                    'text': text, 'features': features})])

    full_summary = ""
    for i, element in enumerate(summaries):
        element[1] = element[1].replace("\n\n", "\n")
        full_summary += f'{element[0]}\n{element[1]}'
        if i < len(summaries) - 1:
            full_summary += '\n\n'

    return full_summary


def generate_document(summaries: str, document_name: str, type_doc: str) -> io.BytesIO:
    """
    Generates a PDF or Word document based on the specified type.
    The generated document is stored in a buffer and returned.

    :param summaries: Summaries of all files
    :param document_name: Document's name
    :param type_doc: Document's type
    :return: io.BytesIO containing the generated document
    """

    if type_doc == ".PDF":
        buffer = io.BytesIO()
        styles = getSampleStyleSheet()
        custom_style = ParagraphStyle(name='CustomStyle', parent=styles['Normal'])
        custom_style.fontSize = 12

        pdf = SimpleDocTemplate(buffer, pagesize=letter, title=document_name)
        summary_list = []

        title = document_name
        title_paragraph = Paragraph(title, styles['Title'])
        title_paragraph.alignment = 1
        summary_list.append(title_paragraph)
        fragments = summaries.strip().split('\n\n')
        for fragment in fragments:
            fragment = fragment.replace('\n\n', '\n')
            sub_fragments = fragment.split('\n')
            for sub_fragment in sub_fragments:
                paragraph = Paragraph(sub_fragment, custom_style)
                summary_list.append(paragraph)
            if fragment != fragments[-1]:
                summary_list.append(Spacer(1, 12))
        pdf.build(summary_list)
        return buffer
    else:
        buffer = io.BytesIO()
        doc = Document()
        title = doc.add_paragraph(document_name)
        title.runs[0].font.bold = True
        title.alignment = 1
        doc.add_paragraph(summaries)
        doc.save(buffer)
        return buffer


def get_youtube_transcript(link: str) -> str:
    """
    Retrieve the transcription of the YouTube video given its link.

    :param link: YouTube video link
    :return: Transcription of the YouTube video
    """
    try:
        loader = YoutubeLoader.from_youtube_url(
            link, add_video_info=True, language=["en", "es"]
        )
        doc_yt = loader.load()
        if doc_yt:
            first_element = doc_yt[0]
            title = first_element.metadata['title']
            author = first_element.metadata['author']
            yt_text = ''
            for page in doc_yt:
                yt_text += page.page_content.replace('\t', ' ')
            yt_transcript = f"Tittle: {title}, Author: {author}, Transcription: {yt_text}"
            return yt_transcript
        else:
            return "fail"
    except Exception as e:
        st.error(str(e))
        st.stop()


def youtube_summarization(yt_transcript: str, llm_name: str, key: str, lang: str) -> str:
    """
    Returns the summary of the YouTube video.
    If the text of the transcription is longer than the model's context, the 'handle_long_text' function is called,
    which will return the summary of that large transcription.
    If the text of the transcription is shorter than the model's context, the video summary will be obtained.

    :param yt_transcript: YouTube video transcription
    :param llm_name: Name of the desired LLM.
    :param key: OpenAI API key.
    :param lang: Chosen language
    :return: Summary of the YouTube video
    """
    matter = "data about a youtube video transcription"
    llm, context_length = llm_choice(llm_name, key, "sum")
    if llm.get_num_tokens(yt_transcript) >= context_length:
        long_response_yt = handle_long_text(llm, context_length, yt_transcript, lang, matter, "")
        return long_response_yt
    else:
        chain = LLMChain(llm=llm, prompt=prompt)
        yt_response = chain.run({'lang': lang, 'matter': matter, 'text': yt_transcript, 'features': ''})
        return yt_response
